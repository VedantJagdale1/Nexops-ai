from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("output/resume/Vedant_Jagdale_Resume_NexOps_AI.docx")
FONT = "Times New Roman"
INK = RGBColor(0, 0, 0)
LINK = RGBColor(5, 99, 193)


def set_run_font(run, size, bold=False, italic=False, color=INK, underline=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.underline = underline


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_bottom_border(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    p_pr.append(borders)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run_element = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    run_properties.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "19")
    run_properties.append(size)
    run_element.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run_element.append(text_node)
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def add_section_header(document, title):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    set_paragraph_spacing(paragraph, before=9, after=3, line=1.0)
    set_bottom_border(paragraph)
    run = paragraph.add_run(title.upper())
    set_run_font(run, 11.6)
    return paragraph


def add_left_right_line(document, left, right, left_size=10.2, right_size=9.7, bold_left=False):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.56), WD_TAB_ALIGNMENT.RIGHT)
    set_paragraph_spacing(paragraph, after=0, line=1.0)
    left_run = paragraph.add_run(left)
    set_run_font(left_run, left_size, bold=bold_left)
    paragraph.add_run("\t")
    right_run = paragraph.add_run(right)
    set_run_font(right_run, right_size)
    return paragraph


def add_subline(document, left, right=None, italic=True):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(7.56), WD_TAB_ALIGNMENT.RIGHT)
    set_paragraph_spacing(paragraph, after=1, line=1.0)
    left_run = paragraph.add_run(left)
    set_run_font(left_run, 9.8, italic=italic)
    if right:
        paragraph.add_run("\t")
        right_run = paragraph.add_run(right)
        set_run_font(right_run, 9.6, italic=italic)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.23)
    paragraph.paragraph_format.first_line_indent = Inches(-0.14)
    paragraph.paragraph_format.keep_together = True
    set_paragraph_spacing(paragraph, after=0.5, line=1.03)
    run = paragraph.add_run(text)
    set_run_font(run, 9.55)
    return paragraph


def add_project(document, name, technologies, bullets):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    set_paragraph_spacing(paragraph, before=4, after=1, line=1.0)
    title = paragraph.add_run(name)
    set_run_font(title, 10.1, bold=True)
    separator = paragraph.add_run(" | ")
    set_run_font(separator, 10.1)
    tech = paragraph.add_run(technologies)
    set_run_font(tech, 9.85, italic=True)
    for bullet in bullets:
        add_bullet(document, bullet)


def add_skill_line(document, label, value):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=0, line=1.0)
    key = paragraph.add_run(f"{label}: ")
    set_run_font(key, 9.55, bold=True)
    detail = paragraph.add_run(value)
    set_run_font(detail, 9.55)
    return paragraph


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.46)
    section.left_margin = Inches(0.46)
    section.right_margin = Inches(0.48)
    section.header_distance = Inches(0.25)
    section.footer_distance = Inches(0.25)

    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.55)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    bullet = document.styles["List Bullet"]
    bullet.base_style = normal
    bullet.font.name = FONT
    bullet._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    bullet.font.size = Pt(9.55)

    if "Resume Section" not in [style.name for style in document.styles]:
        style = document.styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = normal


def build_resume():
    document = Document()
    configure_document(document)

    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name, after=2, line=1.0)
    name_run = name.add_run("VEDANT JAGDALE")
    set_run_font(name_run, 19.0, bold=True)

    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(contact, after=6, line=1.0)
    phone = contact.add_run("+91 8010218846 | ")
    set_run_font(phone, 9.5)
    add_hyperlink(contact, "vedantjagdale04@gmail.com", "mailto:vedantjagdale04@gmail.com")
    separator_one = contact.add_run(" | ")
    set_run_font(separator_one, 9.5)
    add_hyperlink(contact, "LinkedIn", "https://www.linkedin.com")
    separator_two = contact.add_run(" | ")
    set_run_font(separator_two, 9.5)
    add_hyperlink(contact, "GitHub", "https://github.com")

    add_section_header(document, "Education")
    add_left_right_line(document, "Mumbai University", "Aug 2025 - Present", bold_left=True)
    add_subline(document, "Master of Computer Applications (MCA)", "Mumbai, Maharashtra")
    add_bullet(
        document,
        "Relevant Coursework: Data Structures & Algorithms, Advanced Database Management Systems, Mobile Application Development, Software Engineering.",
    )
    add_left_right_line(document, "Mumbai University", "Aug 2022 - Apr 2025", bold_left=True)
    add_subline(document, "Bachelor of Science in Information Technology (B.Sc. IT)", "Mumbai, Maharashtra")

    add_section_header(document, "Experience")
    add_left_right_line(document, "Python Engineer Intern", "June 2024 - May 2025", bold_left=True)
    add_subline(document, "ONLEI Technologies", "Mumbai, India")
    add_bullet(
        document,
        "Engineered 3 scalable REST API endpoints using Flask, reducing integration latency by 20% and supporting high-concurrency data requests for internal web applications.",
    )
    add_bullet(
        document,
        "Optimized deployment cycles by standardizing modular component architecture, improving code reusability across 5+ internal projects and reducing manual debugging time by 10%.",
    )
    add_bullet(
        document,
        "Collaborated with senior engineers to implement automated unit testing, increasing test coverage by 30% and reducing production bug reports by 15%.",
    )

    add_section_header(document, "Projects")
    add_project(
        document,
        "NexOps AI - Multi-Tenant Client Project Management SaaS",
        "React, TypeScript, Node.js, Express, MongoDB, Socket.IO, Docker",
        [
            "Built a multi-tenant SaaS platform for agencies with isolated client, project, Kanban task, ticket, document, invoice, and notification workflows.",
            "Implemented JWT authentication, refresh-token handling, centralized role-based permissions, tenant-scoped data access, Zod validation, and audit logging.",
            "Delivered a responsive dashboard and persistent Kanban board with real-time project chat, presence, typing indicators, task and ticket updates, and notifications via Socket.IO.",
        ],
    )
    add_project(
        document,
        "TradeX - iOS Stock Portfolio Management App",
        "Swift, SwiftUI, SwiftData, MVVM, Yahoo Finance API, Groq LLM API, Apple Charts",
        [
            "Built a high-performance data pipeline using SwiftData and async/await, enabling real-time processing of 2,000+ stocks with 40% faster search response times compared to standard implementations.",
            "Engineered a responsive portfolio dashboard using Apple Charts, maintaining sub-100ms UI latency under high data loads and reducing user time-to-insight by 30% through automated LLM-driven analysis.",
        ],
    )
    add_project(
        document,
        "Dime - iOS Personal Finance App",
        "Swift, SwiftUI, App Intents, WidgetKit, iCloud (CloudKit), Local Authentication",
        [
            "Architected a comprehensive personal finance platform using SwiftUI and iCloud, ensuring 99.9% data consistency across user devices.",
            "Optimized user workflows by integrating App Intents and WidgetKit, reducing expense-logging time by 40% through streamlined home-screen interactions.",
            "Automated recurring transaction management and implemented Face ID security, reducing manual entry effort by 50% while maintaining high-grade data privacy.",
        ],
    )
    add_project(
        document,
        "EquityLens - Stock Analysis & Portfolio Management Web App",
        "Python, Streamlit, TensorFlow (LSTM), Supabase, Plotly",
        [
            "Built a full-stack platform tracking all Nifty 50 stocks with technical indicators, candlestick charts, sector heatmaps, and a 30-day LSTM price-prediction model; backend and database hosted on Supabase.",
        ],
    )
    add_section_header(document, "Technical Skills")
    add_skill_line(document, "Languages", "TypeScript, JavaScript, Python, Java, Swift, C++, SQL")
    add_skill_line(
        document,
        "Frontend",
        "React, Vite, Tailwind CSS, React Router, TanStack Query, Zustand, React Hook Form, Recharts, Axios",
    )
    add_skill_line(
        document,
        "Backend & Data",
        "Node.js, Express.js, MongoDB, Mongoose, Socket.IO, Redis, REST APIs, JWT, bcrypt, Nodemailer",
    )
    add_skill_line(
        document,
        "Testing & DevOps",
        "Zod, Vitest, React Testing Library, Playwright, Docker, GitHub Actions, Swagger/OpenAPI, Git, GitHub",
    )
    add_skill_line(
        document,
        "Additional",
        "SwiftUI, SwiftData, MVVM, Flask, PostgreSQL (Supabase), TensorFlow, Streamlit, Plotly, Agile",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_resume()
